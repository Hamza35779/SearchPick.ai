"""
File parsing service — extracts structured text from uploaded files
so the AI agents can consume them as search context.

Supported:
  - Images   (.jpg, .jpeg, .png, .webp, .bmp, .tiff)
  - CSV      (.csv)
  - Excel    (.xlsx, .xls)
  - Word     (.docx)
"""

import csv
import io
import json
from pathlib import Path
from typing import Any

from fastapi import UploadFile
from PIL import Image
import openpyxl
import docx


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _truncate(text: str, max_chars: int = 4000) -> str:
    """Keep extraction output from blowing up the LLM context window."""
    return text[:max_chars] + "…" if len(text) > max_chars else text


# ─── Individual parsers ───────────────────────────────────────────────────────

def parse_csv(raw: bytes) -> dict[str, Any]:
    text = raw.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    rows = list(reader)
    headers = reader.fieldnames or []
    preview = rows[:25]                          # first 25 rows for context
    return {
        "file_type": "csv",
        "row_count": len(rows),
        "columns": list(headers),
        "preview_rows": preview,
        "summary": (
            f"CSV with {len(rows)} rows and {len(headers)} columns: "
            f"{', '.join(list(headers)[:10])}."
        ),
        "extracted_text": _truncate(
            f"Columns: {', '.join(headers)}\n\n"
            + "\n".join(json.dumps(r) for r in preview)
        ),
    }


def parse_excel(raw: bytes) -> dict[str, Any]:
    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    all_text: list[str] = []
    sheet_summaries: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(c) if c is not None else "" for c in rows[0]]
        data_rows = rows[1:26]                   # first 25 data rows
        sheet_summaries.append(f"Sheet '{sheet_name}': {len(rows)-1} rows, cols: {', '.join(headers[:8])}")
        all_text.append(f"=== Sheet: {sheet_name} ===")
        all_text.append("Columns: " + ", ".join(headers))
        for r in data_rows:
            all_text.append(", ".join(str(c) if c is not None else "" for c in r))

    wb.close()
    extracted = _truncate("\n".join(all_text))
    return {
        "file_type": "excel",
        "sheets": wb.sheetnames,
        "sheet_summaries": sheet_summaries,
        "summary": f"Excel workbook with {len(wb.sheetnames)} sheet(s). " + "; ".join(sheet_summaries),
        "extracted_text": extracted,
    }


def parse_docx(raw: bytes) -> dict[str, Any]:
    doc = docx.Document(io.BytesIO(raw))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Also grab table cell text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    return {
        "file_type": "docx",
        "paragraph_count": len(paragraphs),
        "summary": f"Word document with {len(paragraphs)} paragraphs / {len(full_text)} characters.",
        "extracted_text": _truncate(full_text),
    }


def parse_image(raw: bytes, filename: str) -> dict[str, Any]:
    img = Image.open(io.BytesIO(raw))
    exif_data: dict[str, str] = {}

    # Try to pull EXIF safely
    try:
        raw_exif = img._getexif()  # type: ignore[attr-defined]
        if raw_exif:
            from PIL.ExifTags import TAGS
            exif_data = {
                TAGS.get(k, str(k)): str(v)
                for k, v in raw_exif.items()
                if isinstance(v, (str, int, float, bytes))
            }
    except Exception:
        pass

    # Build a human-readable description for the AI to use as context
    desc_parts = [
        f"Image file: {filename}",
        f"Format: {img.format or 'unknown'}",
        f"Mode: {img.mode}",
        f"Resolution: {img.width}x{img.height} px",
    ]
    if exif_data.get("Make") or exif_data.get("Model"):
        desc_parts.append(f"Camera: {exif_data.get('Make','')} {exif_data.get('Model','')}".strip())
    if exif_data.get("DateTime"):
        desc_parts.append(f"Taken: {exif_data['DateTime']}")

    description = ". ".join(desc_parts)

    return {
        "file_type": "image",
        "width": img.width,
        "height": img.height,
        "format": img.format,
        "exif": exif_data,
        "summary": description,
        "extracted_text": (
            description + "\n\n"
            "[Note: This is an image upload. Use the visual details and any "
            "product-related context the user provides alongside this image "
            "to search for the best matching products.]"
        ),
    }


# ─── Dispatcher ───────────────────────────────────────────────────────────────

async def parse_upload(file: UploadFile) -> dict[str, Any]:
    """
    Reads the uploaded file and routes it to the correct parser.
    Returns a structured dict that the agent graph can consume.
    """
    raw = await file.read()
    suffix = Path(file.filename or "").suffix.lower()

    if suffix == ".csv":
        return parse_csv(raw)
    elif suffix in (".xlsx", ".xls"):
        return parse_excel(raw)
    elif suffix == ".docx":
        return parse_docx(raw)
    elif suffix in (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".tif"):
        return parse_image(raw, file.filename or "image")
    else:
        # Attempt raw UTF-8 text fallback
        try:
            text = raw.decode("utf-8", errors="replace")
            return {
                "file_type": "text",
                "summary": f"Plain text file ({len(text)} characters).",
                "extracted_text": _truncate(text),
            }
        except Exception:
            return {
                "file_type": "unknown",
                "summary": "Unsupported file format.",
                "extracted_text": "",
            }
